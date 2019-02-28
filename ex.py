import re
from docx import Document

filePath = 'docx/AJJ1198_04.04.2016_EMBREAGEM.docx'

def read_file(file_path):
    return Document(file_path)


def read_tables(doc):
    for table in doc.tables:
        for i, row in enumerate(table.rows):
            text = (cell.text for cell in row.cells)

            # Establish the mapping based on the first row
            # headers; these will become the keys of our dictionary
            if i == 0:
                keys = tuple(text)
                continue

            # Construct a dictionary for this row, mapping
            # keys to values for this row
            row_data = dict(zip(keys, text))
            print(row_data)


def read_paragraphs(doc):
    lines = []
    # print(len(doc.paragraphs))
    for line in doc.paragraphs:
        lines.append(line.text)
        # print(line.text)
    return ','.join(lines)


def read_license_plate(doc, lines):
    pattern = '[a-zA-Z]{3}\s*\d{4}$'
    print(lines)
    # print(re.match(pattern, lines))
    # if match:
    #     return match.group()
    # return None


if __name__ == "__main__":
    doc = read_file(filePath)
    lines = read_paragraphs(doc)
    match = read_license_plate(doc, lines)
    print(match)
    # read_tables(doc)





