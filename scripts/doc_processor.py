import os
import re
from docx import Document
from datetime import datetime

# from pydantic import BaseModel
from docx import Document

# from typing import List
from pprint import pprint


SERVICE_DATE = r"(DATA:\s?)([0-9]{1,2}[\.|\/][0-9]{1,2}[\.|\/][0-9]{1,4})"
LICENSE_PLATE = r"(PLACA:\s)([a-zA-Z]{3}\s?[0-9][0-9A-Z][0-9][0-9])"
VEHICLE_NAME = r"(CULO.?:\s?)(.*(?=\s{3})(?<=\S))"
VEHICLE_NUMBER = r"((CHASSI|VIN)\s?:\s?)(\S{1,})"  # match group 3

REGEX_PHONE_A = r"[0-9|(]{1,2}[\s]?[(]?[0-9]{0,3}[)]?[\s]?[0-9]{4,5}[\s]?[-]?[0-9]{4,5}"
REGEX_PHONE_B = r"[0-9]{4,5}[\s]?[-]?[0-9]{4,5}"

CUSTOMER = r"(PROP.?:\s?)(.*)"
KILOMETERS = r"(km\s?)([0-9]{1,3}\.?[0-9]{1,3})"
DOT_SEQUENCE = r"(\.{2,})"


class RegexFieldTypes:
    SERVICE_DATE = r"(DATA:\s?)([0-9]{1,2}[\.|\/][0-9]{1,2}[\.|\/][0-9]{1,4})"
    LICENSE_PLATE = r"(PLACA:\s)([a-zA-Z]{3}\s?[0-9][0-9A-Z][0-9][0-9])"
    VEHICLE_NAME = r"(CULO.?:\s?)(.*(?=\s{3})(?<=\S))"
    VEHICLE_NUMBER = r"((CHASSI|VIN)\s?:\s?)(\S{1,})"  # match group 3
    # PHONE = r'[0-9|(]{1,2}[\s]?[(]?[0-9]{0,3}[)]?[\s]?[0-9]{4,5}[\s]?[-]?[0-9]{4,5}'
    PHONE = r"[0-9]{4,5}[\s]?[-]?[0-9]{4,5}"
    CUSTOMER = r"(PROP.?:\s?)(.*)"
    KILOMETERS = r"(km\s?)([0-9]{1,3}\.?[0-9]{1,3})"
    DOT_SEQUENCE = r"(\.{2,})"


# class ServiceItem(BaseModel):
#     id: int
#     description: str
#     price: float

# class PartItem(BaseModel):
#     id: int
#     code_factory: str
#     description: str
#     genuine: float
#     other: float
#     price: float


class DocProcessor:
    def __init__(self, doc_folder: str, filename: str):
        self._doc_folder = doc_folder
        self._filename = filename
        self._docx = Document("{}{}".format(doc_folder, filename))

    @staticmethod
    def read_table_content(doc_tables):

        content_lines = []
        counter = 1

        for table in doc_tables:
            service_items = []
            # DocProcessor.extractTableData(table) if counter == 1 : DocProcessor.extractParts(table)
            for i, row in enumerate(table.rows):
                text = (cell.text for cell in row.cells)

                # Establish the mapping based on the first row
                # headers; these will become the keys of our dictionary
                if i == 0:
                    keys = tuple(text)
                    continue

                # Construct a dictionary for this row, mapping
                # keys to values for this row
                # pprint(keys)
                row_data = dict(zip(keys, text))

                # print(row_data)
                # for k, v in row_data.items():
                #     item = {}

                #     if 'DESC' in k:
                #         item["descript"]

                if "DESCRIÇÃO" in row_data and "PREÇO" in row_data:
                    if row_data["PREÇO"]:
                        service_items.append(
                            {
                                "type": "service",
                                "description": row_data["DESCRIÇÃO"],
                                "price": row_data["PREÇO"],
                            }
                        )

                if "CÓD. FAB." in row_data and "DISCRIMINAÇÃO" in row_data:
                    if row_data["ORIGINAL"] or row_data["OUTRA"]:
                        service_items.append(
                            {
                                "type": "part",
                                "code_factory": row_data["CÓD. FAB."],
                                "number": 1,
                                "description": row_data["DISCRIMINAÇÃO"],
                                "price": row_data["ORIGINAL"]
                                if row_data["ORIGINAL"]
                                else row_data["OUTRA"],
                            }
                        )

                # print(row_data)

                content_lines.append(row_data)

            counter += 1

            pprint(service_items)
        return content_lines

    @staticmethod
    def remove_dot_sequence(line):
        match = re.search(RegexFieldTypes.DOT_SEQUENCE, line)
        if match:
            return line.replace(match.group(1), "")
        return line

    @staticmethod
    def remove_double_spaces(line):
        return line.replace("  ", " ")

    def read_lines(self):
        service = {}
        lines = []
        others = []
        others_active = False
        index = 0
        for line in self._docx.paragraphs:
            trimmed_text = line.text.strip()
            if trimmed_text != "" and index >= 3:

                text = DocProcessor.remove_dot_sequence(trimmed_text)
                text = DocProcessor.remove_double_spaces(text)
                # print(text)

                if "PLACA" in text:
                    match = re.search(LICENSE_PLATE, text)
                    service["plate"] = match.group(2)

                if "DATA" in text:
                    match = re.search(SERVICE_DATE, text)
                    service["date"] = match.group(2).replace(".", "-")

                if "TEL" in text:
                    result_a = re.findall(REGEX_PHONE_A, text)
                    result_b = re.findall(REGEX_PHONE_B, text)
                    service["phones"] = result_a + result_b

                if "PROP" in text:
                    service["customer"] = text.replace("PROP", "")

                if "PROP" in text:
                    match = re.search(CUSTOMER, text)
                    service["customer"] = match.group(2)

                if "OUTROS" in text:
                    service["others"] = []
                    others_active = True

                if others_active:
                    service["others"].append(text)

                lines.append(text)
            index = index + 1
        # print(service)
        # pprint(lines)
        # exit()
        return lines

    def read_phones(self, lines):
        pattern = RegexFieldTypes.PHONE
        for line in lines:
            match = re.match(pattern, line)
            if match:
                phones = match.group(2).split("/")
                return self.format_phone(phones[0]), self.format_phone(
                    phones[1]
                ) if len(phones) > 1 else self.format_phone(phones[0])
        return "CAMPO_VAZIO", "CAMPO_VAZIO"

    @staticmethod
    def read_field(lines, regex, match_group):
        for line in lines:
            match = re.search(regex, line)
            if match:
                return " ".join(match.group(match_group).strip().split())
        return "CAMPO_VAZIO"

    def format_phone(self, phone_string):
        return phone_string.strip().replace("(", "").replace(")", "").replace(" ", "")

    def format_to_iso_date(self, str_date):
        if str_date == "CAMPO_VAZIO":
            return str_date

        str_date_hyphen = str_date.replace(" ", "").replace(",", ".").replace(".", "-")
        try:
            date_object = datetime.strptime(str_date_hyphen, "%d-%m-%Y").date()
            return date_object.isoformat()
        except Exception as e:
            print(e)
            return "ERRO"

    def write_to_file(self, lines: list):
        f = open("./test.txt", "a")
        for line in lines:
            f.write(line)
        f.close()

    def extract_data(self):
        lines = self.read_lines()

        for line in lines:
            self.write_to_file(line + "\n")

        phones = self.read_phones(lines)

        return {
            "license_plate": DocProcessor.read_field(
                lines, RegexFieldTypes.LICENSE_PLATE, 2
            ),
            "vehicle_name": DocProcessor.read_field(
                lines, RegexFieldTypes.VEHICLE_NAME, 2
            ),
            "vehicle_number": DocProcessor.read_field(
                lines, RegexFieldTypes.VEHICLE_NUMBER, 2
            ),
            "phone_numbers": phones,
            "customer": DocProcessor.read_field(lines, RegexFieldTypes.CUSTOMER, 2),
            "date": self.format_to_iso_date(
                DocProcessor.read_field(lines, RegexFieldTypes.SERVICE_DATE, 2)
            ),
            "tables": DocProcessor.read_table_content(self._docx.tables),
        }
